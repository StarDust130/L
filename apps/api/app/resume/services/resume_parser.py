import zipfile  # 📦 Work with ZIP files
from io import BytesIO  # 🧠 Treat bytes like a file
from pathlib import Path  # 📁 Work with file paths

from app.resume.resume_schema import ExtractedResume  # 📄 Resume response model
from docx import Document  # 📝 Read DOCX files
from fastapi import HTTPException, UploadFile  # 🚨 API errors + uploads
from pypdf import PdfReader  # 📕 Read PDF files

MAX_RESUME_SIZE = 5 * 1024 * 1024  # 📏 Max file size: 5 MB

ALLOWED_EXTENSIONS = {
    ".pdf",  # 📕 PDF allowed
    ".docx",  # 📝 DOCX allowed
    ".txt",  # 📄 TXT allowed
}

# 🔄 Normalize text by removing extra spaces and empty lines
def normalize_text(text: str) -> str:
    # 🔄 Convert different line endings to \n
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    # 🧹 Remove empty lines and extra spaces
    clean_lines = [line.strip() for line in lines if line.strip()]

    # 🔗 Join all clean lines
    return "\n".join(clean_lines)

# 📝 Validate uploaded file for type and size
def validate_file(
    filename: str,
    file_bytes: bytes,
) -> str:
    # 🔍 Get file extension
    extension = Path(filename).suffix.lower()

    # ❌ Check if file type is allowed
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Only PDF, DOCX, and TXT files are allowed.",
        )

    # 📏 Check file size
    if len(file_bytes) > MAX_RESUME_SIZE:
        raise HTTPException(
            status_code=413,
            detail="Resume must be smaller than 5 MB.",
        )

    # 🔎 Check PDF file signature
    if extension == ".pdf" and not file_bytes.startswith(b"%PDF-"):
        raise HTTPException(
            status_code=400,
            detail="This does not look like a valid PDF file.",
        )

    # 📝 Check DOCX is a valid ZIP file
    if extension == ".docx":
        is_zip = zipfile.is_zipfile(BytesIO(file_bytes))

        # ❌ Invalid DOCX file
        if not is_zip:
            raise HTTPException(
                status_code=400,
                detail="This does not look like a valid DOCX file.",
            )

    # ✅ Return valid extension
    return extension

# 📄 .pdf Page Parsing
def parse_pdf(file_bytes: bytes) -> str:
    # 📕 Open PDF from bytes
    reader = PdfReader(BytesIO(file_bytes))

    # 📄 Extract text from every page
    pages = [page.extract_text() or "" for page in reader.pages]

    # 🧹 Clean and combine page text
    return normalize_text("\n".join(pages))

# 📄 .docx Page Parsing
def parse_docx(file_bytes: bytes) -> str:
    # 📝 Open DOCX from bytes
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    # 📄 Read normal paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # 📊 Read tables
    for table in document.tables:
        for row in table.rows:
            # 🔤 Get non-empty cell values
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            # ➕ Add table row as text
            if cells:
                parts.append(" | ".join(cells))

    # 🧹 Clean and combine all text
    return normalize_text("\n".join(parts))

# 📄 .txt Page Parsing
def parse_txt(file_bytes: bytes) -> str:
    try:
        # 📄 Decode TXT using UTF-8
        return normalize_text(
            file_bytes.decode("utf-8-sig"),
        )

    except UnicodeDecodeError as error:
        # ❌ File is not UTF-8
        raise HTTPException(
            status_code=400,
            detail="TXT file must use UTF-8 encoding.",
        ) from error

# 🦣 Extract resume text from uploaded file.
async def extract_resume_text(
    file: UploadFile,
) -> ExtractedResume:
    # 📥 Get uploaded filename
    filename = file.filename or ""

    # 📦 Read uploaded file
    file_bytes = await file.read()

    # ✅ Validate file
    extension = validate_file(
        filename=filename,
        file_bytes=file_bytes,
    )

    # 📕 Parse PDF
    if extension == ".pdf":
        text = parse_pdf(file_bytes)
        file_type = "pdf"

    # 📝 Parse DOCX
    elif extension == ".docx":
        text = parse_docx(file_bytes)
        file_type = "docx"

    # 📄 Parse TXT
    else:
        text = parse_txt(file_bytes)
        file_type = "txt"

    # ❌ Make sure text was extracted
    if not text:
        raise HTTPException(
            status_code=422,
            detail=("No text could be extracted. This may be a scanned PDF."),
        )

    # 📤 Return extracted resume data
    return ExtractedResume(
        filename=filename,
        file_type=file_type,
        text=text,
        character_count=len(text),
    )
